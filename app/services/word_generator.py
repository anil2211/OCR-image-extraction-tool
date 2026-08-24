from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class WordGenerator:

    @staticmethod
    def generate(
        document_data: dict[str, Any],
        output_path: str | Path,
    ) -> Path:

        output_path = Path(
            output_path
        )

        document = Document()

        # ======================================================
        # TITLE
        # ======================================================

        title = document.add_heading(
            "OCR Extracted Document",
            level=1
        )

        title.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        # ======================================================
        # PAGES
        # ======================================================

        pages = document_data.get(
            "pages",
            []
        )

        for page in pages:

            page_number = page.get(
                "page_number",
                1
            )

            document.add_heading(
                f"Page {page_number}",
                level=2
            )

            # --------------------------------------------------
            # TEXT
            # --------------------------------------------------

            text = page.get(
                "text",
                ""
            )

            if text:

                for line in text.splitlines():

                    line = line.strip()

                    if not line:
                        continue

                    paragraph = (
                        document.add_paragraph()
                    )

                    run = paragraph.add_run(
                        line
                    )

                    run.font.size = Pt(
                        11
                    )

            # --------------------------------------------------
            # TABLES
            # --------------------------------------------------

            tables = page.get(
                "tables",
                []
            )

            for table_index, table in enumerate(
                tables,
                start=1
            ):

                rows = (
                    WordGenerator.extract_table_rows(
                        table
                    )
                )

                if not rows:
                    continue

                document.add_paragraph(
                    f"Table {table_index}"
                )

                max_columns = max(
                    len(row)
                    for row in rows
                )

                word_table = (
                    document.add_table(
                        rows=len(rows),
                        cols=max_columns
                    )
                )

                word_table.style = (
                    "Table Grid"
                )

                for row_index, row in enumerate(
                    rows
                ):

                    for column_index, value in enumerate(
                        row
                    ):

                        word_table.cell(
                            row_index,
                            column_index
                        ).text = (
                            ""
                            if value is None
                            else str(value)
                        )

                document.add_paragraph()

        # ======================================================
        # SAVE
        # ======================================================

        document.save(
            output_path
        )

        return output_path

    @staticmethod
    def extract_table_rows(
        table: Any
    ) -> list[list]:

        if table is None:
            return []

        # ------------------------------------------------------
        # Already normalized table
        # ------------------------------------------------------

        if isinstance(
            table,
            list
        ):

            if not table:
                return []

            if all(
                isinstance(
                    row,
                    (list, tuple)
                )
                for row in table
            ):

                return [
                    list(row)
                    for row in table
                ]

        # ------------------------------------------------------
        # Dictionary result
        # ------------------------------------------------------

        if isinstance(
            table,
            dict
        ):

            for key in (
                "structure",
                "res",
                "table",
                "table_res",
                "table_result",
            ):

                value = table.get(
                    key
                )

                if value is None:
                    continue

                nested = (
                    WordGenerator.extract_table_rows(
                        value
                    )
                )

                if nested:
                    return nested

            # --------------------------------------------------
            # HTML table
            # --------------------------------------------------

            html = table.get(
                "html"
            )

            if isinstance(
                html,
                str
            ):

                return (
                    WordGenerator.parse_html_table(
                        html
                    )
                )

            # --------------------------------------------------
            # Cells
            # --------------------------------------------------

            cells = table.get(
                "cells"
            )

            if isinstance(
                cells,
                list
            ):

                rows = []

                for cell in cells:

                    if isinstance(
                        cell,
                        dict
                    ):

                        rows.append([
                            cell.get(
                                "text",
                                ""
                            )
                        ])

                    elif isinstance(
                        cell,
                        str
                    ):

                        rows.append([
                            cell
                        ])

                return rows

        # ------------------------------------------------------
        # String fallback
        # ------------------------------------------------------

        if isinstance(
            table,
            str
        ):

            return [
                [line]
                for line in table.splitlines()
                if line.strip()
            ]

        return []

    @staticmethod
    def parse_html_table(
        html: str
    ) -> list[list]:

        from html.parser import HTMLParser

        class TableParser(
            HTMLParser
        ):

            def __init__(self):

                super().__init__()

                self.rows = []

                self.current_row = None

                self.current_cell = None

            def handle_starttag(
                self,
                tag,
                attrs
            ):

                tag = tag.lower()

                if tag == "tr":

                    self.current_row = []

                elif tag in (
                    "td",
                    "th"
                ):

                    self.current_cell = ""

            def handle_data(
                self,
                data
            ):

                if self.current_cell is not None:

                    self.current_cell += data

            def handle_endtag(
                self,
                tag
            ):

                tag = tag.lower()

                if tag in (
                    "td",
                    "th"
                ):

                    if self.current_row is not None:

                        self.current_row.append(
                            self.current_cell.strip()
                        )

                    self.current_cell = None

                elif tag == "tr":

                    if (
                        self.current_row
                        and any(
                            str(value).strip()
                            for value
                            in self.current_row
                        )
                    ):

                        self.rows.append(
                            self.current_row
                        )

                    self.current_row = None

        parser = TableParser()

        parser.feed(
            html
        )

        return parser.rows